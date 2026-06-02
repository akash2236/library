import os
import re
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    """Sets background shading color for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    """Adds internal padding/margins inside a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_code_callout(doc, code_text):
    """Creates a beautifully shaded single-cell table to serve as a code callout box."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    
    # Premium light-gray surface backdrop shading
    set_cell_background(cell, "F3F4F6")
    set_cell_margins(cell, 140, 140, 200, 200)
    
    # Style the text inside the cell as monospace
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    
    lines = code_text.strip().split('\n')
    for idx, line in enumerate(lines):
        if idx > 0:
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9.0)
        run.font.color.rgb = RGBColor(31, 41, 55) # Sleek charcoal gray

def parse_runs_into_paragraph(paragraph, line_text, is_bullet=False):
    """Parses markdown inline bold tags (**text**) into styled runs inside a paragraph."""
    # Matches markdown links [text](url) and strip them to raw text for simplicity
    line_text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', line_text)
    
    # Split on markdown bold markers
    parts = line_text.split('**')
    for index, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.font.name = 'Calibri'
        run.font.size = Pt(11.0 if not is_bullet else 10.5)
        run.font.color.rgb = RGBColor(55, 65, 81) # Slate gray readability
        
        # Every odd index inside a split represents bold text
        if index % 2 == 1:
            run.bold = True
            run.font.color.rgb = RGBColor(17, 24, 39) # Deep dark gray

def main():
    md_path = "technical_documentation.md"
    docx_path = "technical_documentation.docx"
    
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return
        
    print(f"Reading documentation from {md_path}...")
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    doc = docx.Document()
    
    # Configure formal page layout margins
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)
        
    # State tracking variables
    in_code_block = False
    code_accumulator = []
    
    in_table = False
    table_rows = []
    
    for line in lines:
        stripped = line.strip()
        
        # --- 1. Handle Code Block Toggles ---
        if stripped.startswith("```"):
            if in_code_block:
                # Close and render the code block
                add_code_callout(doc, "\n".join(code_accumulator))
                code_accumulator = []
                in_code_block = False
            else:
                in_code_block = True
            continue
            
        if in_code_block:
            code_accumulator.append(line.rstrip('\n'))
            continue
            
        # --- 2. Handle Markdown Tables ---
        if stripped.startswith("|"):
            in_table = True
            # Parse row columns and filter empty bounds
            cols = [col.strip() for col in stripped.split("|")[1:-1]]
            
            # Check if this row is just table separation headers (e.g. | :--- | :--- |)
            if all(re.match(r'^:?-+:?$', c) for c in cols):
                continue
            table_rows.append(cols)
            continue
        else:
            if in_table and table_rows:
                # Close and render accumulated table
                num_cols = len(table_rows[0])
                table = doc.add_table(rows=len(table_rows), cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                for r_idx, row_data in enumerate(table_rows):
                    row = table.rows[r_idx]
                    
                    # Highlight header rows in deep violet gradient fill
                    is_header = (r_idx == 0)
                    bg_color = "6D28D9" if is_header else ("F9FAFB" if r_idx % 2 == 1 else "FFFFFF")
                    
                    for c_idx, cell_content in enumerate(row_data):
                        # Ensure cell index is within bounds
                        if c_idx < len(row.cells):
                            cell = row.cells[c_idx]
                            cell.text = ""
                            set_cell_background(cell, bg_color)
                            set_cell_margins(cell, 100, 100, 140, 140)
                            
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_before = Pt(3)
                            p.paragraph_format.space_after = Pt(3)
                            p.paragraph_format.line_spacing = 1.15
                            
                            # Standardize text variables
                            run = p.add_run(cell_content)
                            run.font.name = 'Calibri'
                            run.font.size = Pt(10.5 if is_header else 10.0)
                            
                            if is_header:
                                run.bold = True
                                run.font.color.rgb = RGBColor(255, 255, 255) # White text
                            else:
                                run.font.color.rgb = RGBColor(31, 41, 55)
                                
                doc.add_paragraph().paragraph_format.space_after = Pt(6)
                table_rows = []
                in_table = False
            
        # Ignore empty lines
        if not stripped:
            continue
            
        # --- 3. Parse Headings ---
        if stripped.startswith("# "):
            heading_text = stripped[2:]
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(8)
            h.paragraph_format.keep_with_next = True
            run = h.add_run(heading_text)
            run.bold = True
            run.font.name = 'Georgia'
            run.font.size = Pt(22.0)
            run.font.color.rgb = RGBColor(76, 29, 149) # Deep Purple Title
            continue
            
        if stripped.startswith("## "):
            heading_text = stripped[3:]
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(6)
            h.paragraph_format.keep_with_next = True
            run = h.add_run(heading_text)
            run.bold = True
            run.font.name = 'Georgia'
            run.font.size = Pt(16.0)
            run.font.color.rgb = RGBColor(109, 40, 217) # Medium Purple Header
            continue
            
        if stripped.startswith("### "):
            heading_text = stripped[4:]
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(4)
            h.paragraph_format.keep_with_next = True
            run = h.add_run(heading_text)
            run.bold = True
            run.font.name = 'Georgia'
            run.font.size = Pt(12.5)
            run.font.color.rgb = RGBColor(139, 92, 246) # Light Violet Sub-header
            continue
            
        # --- 4. Parse Lists & Bullet Points ---
        if stripped.startswith("* ") or stripped.startswith("- "):
            list_content = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            parse_runs_into_paragraph(p, list_content, is_bullet=True)
            continue
            
        if stripped.startswith("1. ") or re.match(r'^\d+\.\s', stripped):
            list_content = re.sub(r'^\d+\.\s+', '', stripped)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            parse_runs_into_paragraph(p, list_content, is_bullet=True)
            continue
            
        # Indented bullets inside modules specifications (e.g. * *3.1 Button Component*)
        if stripped.startswith("   * ") or stripped.startswith("  * "):
            list_content = re.sub(r'^\s+\*\s+', '', stripped)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            parse_runs_into_paragraph(p, list_content, is_bullet=True)
            continue
            
        # Ignore structural Markdown page break lines
        if stripped == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            p_border = p.add_run("______________________________________________________________________")
            p_border.font.color.rgb = RGBColor(209, 213, 219) # Light gray divider line
            continue
            
        # --- 5. Parse Normal Paragraph Text ---
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        parse_runs_into_paragraph(p, stripped)
        
    print(f"Compilation finished. Saving binary Word file to {docx_path}...")
    doc.save(docx_path)
    print("Word Document compiled successfully!")

if __name__ == "__main__":
    main()
